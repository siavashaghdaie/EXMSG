#!/usr/bin/env python3
"""
Linda Document Generator — Creates professional .docx and .pdf files
from markdown-style text content.

Usage:
  python3 generate_document.py --format docx --output path.docx --title "Title" --content "markdown text"
  python3 generate_document.py --format pdf  --output path.pdf  --title "Title" --content "markdown text"
  python3 generate_document.py --format docx --output path.docx --title "Title" --content-file /tmp/content.txt
"""

import argparse
import re
import sys
import os


# ─── DOCX Generation ────────────────────────────────────────────────────────

def generate_docx(content: str, title: str, output_path: str):
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    doc = Document()

    # Page setup - US Letter with 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Configure default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Configure heading styles
    for level, (size, color, space_before, space_after) in enumerate([
        (26, RGBColor(0x1A, 0x56, 0x8E), 24, 12),   # Heading 1
        (20, RGBColor(0x2E, 0x74, 0xB5), 18, 8),     # Heading 2
        (16, RGBColor(0x1F, 0x4D, 0x78), 14, 6),     # Heading 3
    ], start=1):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Calibri'
        hstyle.font.size = Pt(size)
        hstyle.font.color.rgb = color
        hstyle.font.bold = True
        hstyle.paragraph_format.space_before = Pt(space_before)
        hstyle.paragraph_format.space_after = Pt(space_after)
        hstyle.paragraph_format.keep_with_next = True

    # Add title
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_after = Pt(4)
        run = title_para.add_run(title)
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
        run.font.bold = True
        run.font.name = 'Calibri'

        # Subtle separator line
        line_para = doc.add_paragraph()
        line_para.paragraph_format.space_after = Pt(16)
        pBdr = line_para._p.get_or_add_pPr()
        bottom_border = pBdr.makeelement(qn('w:pBdr'), {})
        bottom = bottom_border.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '6',
            qn('w:space'): '1',
            qn('w:color'): '2E74B5',
        })
        bottom_border.append(bottom)
        pBdr.append(bottom_border)

    # Parse and add content
    lines = content.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Table detection (markdown tables with | separators)
        if '|' in stripped and stripped.strip().startswith('|'):
            # Collect all table rows
            table_rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1

            if table_rows:
                _add_table(doc, table_rows)
            continue

        # Headings
        if stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_heading(level=3)
            _add_formatted_runs(p, text)
        elif stripped.startswith('## '):
            text = stripped[3:]
            p = doc.add_heading(level=2)
            _add_formatted_runs(p, text)
        elif stripped.startswith('# '):
            text = stripped[2:]
            p = doc.add_heading(level=1)
            _add_formatted_runs(p, text)
        # Bullet list
        elif re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, text)
        # Numbered list
        elif re.match(r'^\d+[.)]\s+', stripped):
            text = re.sub(r'^\d+[.)]\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_runs(p, text)
        # Horizontal rule
        elif re.match(r'^[-*_]{3,}$', stripped):
            line_para = doc.add_paragraph()
            line_para.paragraph_format.space_before = Pt(8)
            line_para.paragraph_format.space_after = Pt(8)
            pBdr = line_para._p.get_or_add_pPr()
            bottom_border = pBdr.makeelement(qn('w:pBdr'), {})
            bottom = bottom_border.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC',
            })
            bottom_border.append(bottom)
            pBdr.append(bottom_border)
        # Blockquote
        elif stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add left border
            pBdr = p._p.get_or_add_pPr()
            border_el = pBdr.makeelement(qn('w:pBdr'), {})
            left = border_el.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '12',
                qn('w:space'): '8',
                qn('w:color'): '2E74B5',
            })
            border_el.append(left)
            pBdr.append(border_el)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)

        i += 1

    doc.save(output_path)


def _add_formatted_runs(paragraph, text: str):
    """Parse inline markdown formatting (**bold**, *italic*) into Word runs."""
    from docx.shared import RGBColor
    # Pattern: **bold**, *italic*, `code`, plain text
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))'
    for match in re.finditer(pattern, text):
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # Code
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif match.group(5):  # Plain
            paragraph.add_run(match.group(5))


def _add_table(doc, rows):
    """Add a formatted table to the document."""
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style table
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                _add_formatted_runs(p, cell_text)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

                # Header row styling
                if row_idx == 0:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # Blue background for header
                    shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {
                        qn('w:val'): 'clear',
                        qn('w:color'): 'auto',
                        qn('w:fill'): '2E74B5',
                    })
                    cell._tc.get_or_add_tcPr().append(shading)
                else:
                    # Alternate row shading
                    if row_idx % 2 == 0:
                        shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {
                            qn('w:val'): 'clear',
                            qn('w:color'): 'auto',
                            qn('w:fill'): 'F2F7FB',
                        })
                        cell._tc.get_or_add_tcPr().append(shading)

    # Add spacing after table
    doc.add_paragraph()


# ─── PDF Generation ──────────────────────────────────────────────────────────

def generate_pdf(content: str, title: str, output_path: str):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, ListFlowable, ListItem, PageBreak
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        topMargin=1*inch,
        bottomMargin=1*inch,
        leftMargin=1*inch,
        rightMargin=1*inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        'DocTitle',
        parent=styles['Title'],
        fontSize=24,
        textColor=HexColor('#1A568E'),
        spaceAfter=4,
        fontName='Helvetica-Bold',
    ))
    styles.add(ParagraphStyle(
        'DocHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=HexColor('#2E74B5'),
        spaceBefore=20,
        spaceAfter=8,
        fontName='Helvetica-Bold',
    ))
    styles.add(ParagraphStyle(
        'DocHeading2',
        parent=styles['Heading2'],
        fontSize=15,
        textColor=HexColor('#2E74B5'),
        spaceBefore=16,
        spaceAfter=6,
        fontName='Helvetica-Bold',
    ))
    styles.add(ParagraphStyle(
        'DocHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=HexColor('#1F4D78'),
        spaceBefore=12,
        spaceAfter=4,
        fontName='Helvetica-Bold',
    ))
    styles.add(ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        textColor=HexColor('#333333'),
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        fontName='Helvetica',
    ))
    styles.add(ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        textColor=HexColor('#333333'),
        leftIndent=24,
        spaceAfter=3,
        fontName='Helvetica',
    ))
    styles.add(ParagraphStyle(
        'DocQuote',
        parent=styles['Normal'],
        fontSize=11,
        leading=15,
        textColor=HexColor('#555555'),
        leftIndent=30,
        fontName='Helvetica-Oblique',
        spaceAfter=8,
        spaceBefore=8,
        borderPadding=6,
    ))

    elements = []

    # Title
    if title:
        elements.append(Paragraph(title, styles['DocTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=2, color=HexColor('#2E74B5'),
            spaceBefore=2, spaceAfter=16,
        ))

    # Parse content
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if not stripped:
            elements.append(Spacer(1, 6))
            i += 1
            continue

        # Table
        if '|' in stripped and stripped.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                if re.match(r'^\|[\s\-:|]+\|$', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1

            if table_rows:
                _add_pdf_table(elements, table_rows, styles)
            continue

        # Headings
        if stripped.startswith('### '):
            text = _md_to_rl(stripped[4:])
            elements.append(Paragraph(text, styles['DocHeading3']))
        elif stripped.startswith('## '):
            text = _md_to_rl(stripped[3:])
            elements.append(Paragraph(text, styles['DocHeading2']))
        elif stripped.startswith('# '):
            text = _md_to_rl(stripped[2:])
            elements.append(Paragraph(text, styles['DocHeading1']))
        # Bullet
        elif re.match(r'^[-*]\s+', stripped):
            text = _md_to_rl(re.sub(r'^[-*]\s+', '', stripped))
            elements.append(Paragraph(f'\u2022  {text}', styles['DocBullet']))
        # Numbered list
        elif re.match(r'^(\d+)[.)]\s+', stripped):
            num = re.match(r'^(\d+)', stripped).group(1)
            text = _md_to_rl(re.sub(r'^\d+[.)]\s+', '', stripped))
            elements.append(Paragraph(f'{num}.  {text}', styles['DocBullet']))
        # Horizontal rule
        elif re.match(r'^[-*_]{3,}$', stripped):
            elements.append(HRFlowable(
                width="100%", thickness=1, color=HexColor('#CCCCCC'),
                spaceBefore=8, spaceAfter=8,
            ))
        # Blockquote
        elif stripped.startswith('> '):
            text = _md_to_rl(stripped[2:])
            elements.append(Paragraph(text, styles['DocQuote']))
        # Normal
        else:
            text = _md_to_rl(stripped)
            elements.append(Paragraph(text, styles['DocBody']))

        i += 1

    doc.build(elements)


def _md_to_rl(text: str) -> str:
    """Convert markdown inline formatting to ReportLab XML tags."""
    # Escape XML special chars first
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Code
    text = re.sub(r'`(.+?)`', r'<font name="Courier" color="#C7254E">\1</font>', text)
    return text


def _add_pdf_table(elements, rows, styles):
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import Table, TableStyle, Spacer, Paragraph

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    # Normalize rows
    data = []
    for row in rows:
        padded = row + [''] * (num_cols - len(row))
        data.append([Paragraph(_md_to_rl(c), styles['DocBody']) for c in padded])

    col_width = 6.5 * inch / num_cols
    table = Table(data, colWidths=[col_width] * num_cols)

    style_commands = [
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#2E74B5')),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#CCCCCC')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]

    # Alternate row shading
    for row_idx in range(1, len(data)):
        if row_idx % 2 == 0:
            style_commands.append(('BACKGROUND', (0, row_idx), (-1, row_idx), HexColor('#F2F7FB')))

    table.setStyle(TableStyle(style_commands))
    elements.append(Spacer(1, 4))
    elements.append(table)
    elements.append(Spacer(1, 8))


# ─── Main ────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate DOCX or PDF document')
    parser.add_argument('--format', required=True, choices=['docx', 'pdf'])
    parser.add_argument('--output', required=True)
    parser.add_argument('--title', default='')
    parser.add_argument('--content', default='')
    parser.add_argument('--content-file', default='')
    args = parser.parse_args()

    content = args.content
    if args.content_file:
        with open(args.content_file, 'r', encoding='utf-8') as f:
            content = f.read()

    if not content:
        print('Error: No content provided', file=sys.stderr)
        sys.exit(1)

    try:
        if args.format == 'docx':
            generate_docx(content, args.title, args.output)
        elif args.format == 'pdf':
            generate_pdf(content, args.title, args.output)
        print(f'OK:{os.path.getsize(args.output)}')
    except Exception as e:
        print(f'Error: {e}', file=sys.stderr)
        sys.exit(1)
